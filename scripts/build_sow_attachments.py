"""SOW v1.1 별첨 docx 3종 일괄 생성 (창업지원단 제출용).

생성물
- docs/Pathwave_MVP_FunctionSpec_v1.0.docx   — 별첨 1 (기능정의서)
- docs/Pathwave_MVP_BeaconProtocol_v1.0.docx — 별첨 2 (비콘 클라우드 핸드셰이크)
- docs/Pathwave_MVP_NDA_v1.0.docx            — 비밀유지서약 (NDA)

톤
- 캐주얼·보수적. 창업지원단 과제 수행 완료 안전 우선.
- 정교한 스펙은 docs/internal/spec/*.md (실제 개발 트랙) 에서.

폰트
- SOW v1.1 과 동일하게 "맑은 고딕" 통일.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# ════════════════════════════════════════════════════════════════════
# 공용 헬퍼 — 한글 폰트 ("맑은 고딕") + 표 / 제목 / 본문
# ════════════════════════════════════════════════════════════════════
FONT = '맑은 고딕'


def set_font(run, name=FONT, size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 한글(eastAsia) 폰트도 명시 — Word 가 동아시아 텍스트에 이걸 사용
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


def add_title(doc, text):
    """문서 제목 — 굵게, 14pt, 가운데."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=18, bold=True)
    return p


def add_h1(doc, text):
    """1단계 헤딩 — 12pt 굵게."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=13, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, size=10)
    return p


def add_table(doc, rows):
    """rows = [[c1, c2, ...], ...]. 첫 행 = 헤더 (굵게).
    셀 안 모든 텍스트는 맑은 고딕 10pt.
    """
    if not rows:
        return None
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10, bold=(ri == 0))
    return t


def set_default_font(doc):
    """기본 스타일 폰트를 '맑은 고딕' 으로."""
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:cs'), FONT)


# ════════════════════════════════════════════════════════════════════
# 별첨 1 — Pathwave 기능정의서 v1.0
# ════════════════════════════════════════════════════════════════════
def build_function_spec(out_path: Path):
    doc = Document()
    set_default_font(doc)

    add_title(doc, 'Pathwave 기능정의서 v1.0')
    add_para(doc, '(과업지시서 별첨 1 · 창업지원단 제출용)')
    add_para(doc, '작성: 주식회사 트리거소프트 · 2026.05.26 · v1.0')
    add_para(doc, '')

    add_h1(doc, '1. 개요')
    add_para(doc,
        '본 문서는 Pathwave MVP (1단계 - BLE 기반 WiFi 자동연결 핵심 기능) 의 '
        '기능을 사용자 앱, 시설관리자(SP) 웹, 관리자(슈퍼어드민) 웹의 3개 콘솔 단위로 정의한다. '
        '본 문서와 과업지시서가 상충할 경우 과업지시서를 우선한다.')

    add_h1(doc, '2. 시스템 구성')
    add_para(doc, '3 콘솔 + 1 백엔드 구조:')
    add_bullet(doc, '사용자 앱 (iOS / Android, Flutter) — 사용자 측 인터페이스')
    add_bullet(doc, '시설관리자(SP) 웹 (반응형 PC 웹) — 매장 사장 측 인터페이스')
    add_bullet(doc, '관리자(슈퍼어드민) 웹 (PC 웹) — 발주처 운영자 측 인터페이스')
    add_bullet(doc, '백엔드 (Flask + SQLite) — JWT 인증, 권한 관리, 데이터 게이트웨이')
    add_para(doc, '데이터는 발주처가 보유한 마스터 DB 1개에서 게이트웨이 형태로 제공된다.')

    add_h1(doc, '3. 사용자 앱 기능')
    add_h2(doc, '3.1 회원가입 / 로그인')
    add_bullet(doc, '이메일 + 비밀번호 가입 (소셜 로그인은 본 단계 OOS)')
    add_bullet(doc, 'JWT 토큰 기반 세션 유지')
    add_bullet(doc, '약관 동의 (이용약관 / 개인정보처리방침 / 위치기반 약관 등)')
    add_h2(doc, '3.2 BLE 비콘 인식 및 WiFi 자동연결')
    add_bullet(doc, 'BLE 5.x 비콘 광고 패킷 스캔 (포그라운드)')
    add_bullet(doc, '클라우드 검증 후 WiFi 자동연결 (별첨 2 비콘 프로토콜 참조)')
    add_bullet(doc, '1회 확인 후 동일 비콘 영역 내 재진입 시 자동 재연결')
    add_h2(doc, '3.3 매장 정보')
    add_bullet(doc, '주변 매장 목록 (지도/리스트)')
    add_bullet(doc, '매장 상세 (영업시간, 공지, 메뉴)')
    add_h2(doc, '3.4 스탬프 / 쿠폰')
    add_bullet(doc, '방문 스탬프 자동 적립 (비콘 인식 기반)')
    add_bullet(doc, '쿠폰 수신 / 사용 (쿠폰 기능은 일정 협의에 따라 P2 이관 가능)')
    add_h2(doc, '3.5 채팅 / 푸시 알림')
    add_bullet(doc, '매장-사용자 1:1 채팅 (텍스트 + 이미지)')
    add_bullet(doc, 'FCM 푸시 알림 (공지 / 쿠폰 / 채팅)')

    add_h1(doc, '4. 시설관리자(SP) 웹 기능')
    add_h2(doc, '4.1 회원가입 / 매장 등록')
    add_bullet(doc, '사업자 정보 입력 + 카테고리 선택 (자유입력 금지, DB 목록만)')
    add_bullet(doc, '1 계정 = 1 매장 정책 (다중 매장은 본 단계 OOS)')
    add_bullet(doc, '슈퍼어드민 승인 후 활성화')
    add_h2(doc, '4.2 비콘 페어링')
    add_bullet(doc, '슈퍼어드민이 발급한 비콘 ID 를 매장에 claim')
    add_bullet(doc, '신호 강도 임계값 설정 (사전 정의값 기반, 미세 조정 가능)')
    add_h2(doc, '4.3 캠페인 설정')
    add_bullet(doc, '스탬프 / 쿠폰 발급 캠페인 등록')
    add_bullet(doc, '공지사항 / 영업시간 / 메뉴 관리')
    add_h2(doc, '4.4 결제 / 정산')
    add_bullet(doc, '결제수단 등록 (토스페이먼츠 연동)')
    add_bullet(doc, '이용현황 / 정산내역 조회')

    add_h1(doc, '5. 관리자(슈퍼어드민) 웹 기능')
    add_bullet(doc, 'SP 가입 심사 / 승인 / 거절')
    add_bullet(doc, '비콘 자산 등록 (CSV 입고) / 회수 / 이력 관리')
    add_bullet(doc, '카테고리 마스터 관리 (추가 / 수정 / 비활성화)')
    add_bullet(doc, '이용 통계 대시보드 (기본)')
    add_bullet(doc, '정산 처리 / 신고·문의 처리')
    add_bullet(doc, '공지 / 약관 마스터 관리')

    add_h1(doc, '6. 비콘 프로토콜')
    add_para(doc, '상세는 본 SOW 별첨 2 "비콘 클라우드 핸드셰이크 프로토콜 명세" 를 따른다.')

    add_h1(doc, '7. 본 단계 OOS (Out of Scope)')
    add_bullet(doc, '소셜 로그인 (카카오 / 구글 / 애플)')
    add_bullet(doc, '다중 매장 (1 계정 다 매장)')
    add_bullet(doc, '쿠폰 기능 (업무 일정 및 진행 상황에 따라 P2 로 이관 가능)')
    add_bullet(doc, '회사 및 서비스 홍보 웹사이트 (별도 외주 발송 — 별도 과업지시서 필요)')
    add_bullet(doc, '개시 후 유지보수 (본 과업의 하자보수 기간 이후는 별도 계약)')

    add_para(doc, '')
    add_para(doc, '─── 문서 끝 ───')

    doc.save(out_path)
    print(f'  ✓ {out_path.name}')


# ════════════════════════════════════════════════════════════════════
# 별첨 2 — 비콘 클라우드 핸드셰이크 프로토콜 v1.0
# ════════════════════════════════════════════════════════════════════
def build_beacon_protocol(out_path: Path):
    doc = Document()
    set_default_font(doc)

    add_title(doc, '비콘 클라우드 핸드셰이크 프로토콜 명세 v1.0')
    add_para(doc, '(과업지시서 별첨 2 · 창업지원단 제출용)')
    add_para(doc, '작성: 주식회사 트리거소프트 · 2026.05.26 · v1.0')
    add_para(doc, '')

    add_h1(doc, '1. 개요')
    add_para(doc,
        '본 문서는 Pathwave MVP 1단계의 BLE 비콘 ↔ 사용자 앱 ↔ 클라우드 간 '
        '핸드셰이크 프로토콜의 표준 절차를 정의한다. '
        '본 단계는 테스트용 비콘 9대를 기준으로 검증한다. '
        '양산 단계의 추가 발주 및 운영 오류 시 디버깅 지원은 별도 협의로 진행한다.')

    add_h1(doc, '2. 구성 요소')
    add_table(doc, [
        ['구성', '설명'],
        ['비콘 (BLE 5.x)', 'UUID 광고 + 신호 만료 (TTL) + nonce. 양산 OTA 가능 구조.'],
        ['사용자 앱', 'BLE 스캔 → 광고 패킷 수신 → 클라우드 검증 요청'],
        ['클라우드 (백엔드)', '비콘 검증 + WiFi 인증 토큰 발급'],
        ['WiFi AP', '인증 토큰 기반 사용자 접속 허용'],
    ])

    add_h1(doc, '3. 광고 패킷 구조 (개요)')
    add_bullet(doc, 'UUID — 비콘 식별자 (암호화 처리)')
    add_bullet(doc, 'Major / Minor — 매장 / 위치 식별')
    add_bullet(doc, 'Nonce — 매 광고 주기마다 갱신 (재사용 방지)')
    add_bullet(doc, 'TTL — 광고 유효 시간 (재사용 검증)')
    add_para(doc, '※ 구체 비트 단위 명세는 양산 단계 협의 시 확정한다.')

    add_h1(doc, '4. 핸드셰이크 흐름')
    add_para(doc, '아래 순서로 진행된다.')
    add_bullet(doc, '1) 사용자 앱이 BLE 스캔으로 비콘 광고 패킷 수신.')
    add_bullet(doc, '2) UUID + Nonce + TTL 을 클라우드에 검증 요청.')
    add_bullet(doc, '3) 클라우드가 비콘 마스터 DB 와 대조해 검증 → 인증 토큰 발급.')
    add_bullet(doc, '4) 사용자 앱이 인증 토큰을 사용해 WiFi AP 에 접속.')
    add_bullet(doc, '5) 사용자 앱이 동일 매장 비콘을 재인식한 경우 캐시된 토큰으로 즉시 재연결.')

    add_h1(doc, '5. 보안 고려사항')
    add_bullet(doc, 'Nonce 재사용 거부 (Replay 방지)')
    add_bullet(doc, 'TTL 만료 시 거부')
    add_bullet(doc, '비콘 UUID 는 광고 시 암호화 처리')
    add_bullet(doc, '인증 토큰은 짧은 유효시간 + 갱신 토큰 구조')
    add_para(doc, '※ 상세 위협 모델 및 대응은 양산 단계 협의 시 확정한다.')

    add_h1(doc, '6. OTA 업데이트')
    add_bullet(doc, '비콘 펌웨어는 OTA 업데이트가 가능한 구조로 설계된다.')
    add_bullet(doc, '본 단계 (테스트용 9대) 는 수동 펌웨어 적용 가능.')
    add_bullet(doc, '양산 단계의 대규모 OTA 운영은 별도 협의로 진행한다.')

    add_h1(doc, '7. 본 단계 적용 범위')
    add_bullet(doc, '테스트용 비콘 9대 기준으로 위 절차를 검증한다.')
    add_bullet(doc, '추가 발주 및 운영 오류 시 디버깅 지원은 별도 협의로 진행한다.')
    add_bullet(doc, 'KC 인증은 정식 양산 단계에서 진행한다.')

    add_para(doc, '')
    add_para(doc, '─── 문서 끝 ───')

    doc.save(out_path)
    print(f'  ✓ {out_path.name}')


# ════════════════════════════════════════════════════════════════════
# 별첨 — 비밀유지서약서 (NDA) v1.0
# ════════════════════════════════════════════════════════════════════
def build_nda(out_path: Path):
    doc = Document()
    set_default_font(doc)

    add_title(doc, '비밀유지서약서 (NDA) v1.0')
    add_para(doc, '주식회사 트리거소프트 · 2026.05.26 · v1.0')
    add_para(doc, '')

    add_h1(doc, '제 1 조 (목적)')
    add_para(doc,
        '본 서약서는 주식회사 트리거소프트(이하 "발주처") 와 본 과업의 수행사(이하 "수행사") '
        '간의 Pathwave MVP 개발 용역과 관련하여 양 당사자가 공유하는 정보의 비밀유지에 관한 사항을 규정함을 목적으로 한다.')

    add_h1(doc, '제 2 조 (비밀정보의 정의)')
    add_para(doc,
        '본 서약서에서 "비밀정보" 라 함은 본 과업 수행 과정에서 일방이 타방에게 제공하거나 '
        '취득한 모든 기술적·사업적·재무적 정보로서, 다음 각 호를 포함하되 이에 한정되지 아니한다.')
    add_bullet(doc, '소스코드, 디자인 원본, 비콘 프로토콜 명세, API 명세')
    add_bullet(doc, '사용자 / 매장 데이터 (실제 데이터 및 테스트 데이터)')
    add_bullet(doc, '발주처의 사업 계획, 수익 모델, 회계 정보')
    add_bullet(doc, '본 과업의 결과로 생성되는 모든 산출물 및 파생물')

    add_h1(doc, '제 3 조 (비밀유지 의무)')
    add_para(doc, '양 당사자는 다음 의무를 부담한다.')
    add_bullet(doc, '비밀정보를 본 과업 수행 목적 외 다른 목적으로 사용하지 아니한다.')
    add_bullet(doc, '비밀정보를 사전 서면 동의 없이 제3자에게 누설 또는 공개하지 아니한다.')
    add_bullet(doc, '비밀정보의 분실·도난·유출을 방지하기 위한 합리적 보안 조치를 취한다.')
    add_bullet(doc, '본 과업 종료 또는 본 서약서 효력 종료 시, 상대방의 요청에 따라 비밀정보를 반환 또는 파기한다.')

    add_h1(doc, '제 4 조 (예외)')
    add_para(doc, '다음 각 호의 정보는 비밀정보에서 제외된다.')
    add_bullet(doc, '공개된 정보 또는 공지의 사실')
    add_bullet(doc, '상대방으로부터 제공받기 이전에 이미 보유하고 있던 정보')
    add_bullet(doc, '법령 또는 정부기관의 요구에 따라 공개되어야 하는 정보 (단, 이 경우 상대방에게 즉시 통지한다)')

    add_h1(doc, '제 5 조 (효력 기간)')
    add_para(doc,
        '본 서약서의 효력은 양 당사자가 서명한 날로부터 발생하며, 본 과업의 종료 후에도 양 당사자의 '
        '합의로 정한 기간 동안 유지된다. 합의 기간은 본 과업 계약서 또는 별도 합의서에서 정한다.')

    add_h1(doc, '제 6 조 (위반에 대한 책임)')
    add_para(doc,
        '본 서약서를 위반한 당사자는 상대방에게 발생한 손해를 배상할 책임을 진다. '
        '단, 손해배상의 범위는 양 당사자의 합의 또는 본 과업 계약서가 정한 바에 따른다.')

    add_h1(doc, '제 7 조 (관할)')
    add_para(doc, '본 서약서와 관련된 분쟁은 대한민국 법령에 따르며, 발주처 본점 소재지 관할 법원을 1심 관할 법원으로 한다.')

    add_para(doc, '')
    add_h1(doc, '서 명')
    add_para(doc, '발주처: 주식회사 트리거소프트')
    add_para(doc, '대표이사: ________________________ (인)')
    add_para(doc, '일자: ________________')
    add_para(doc, '')
    add_para(doc, '수행사: ________________________')
    add_para(doc, '대표: ________________________ (인)')
    add_para(doc, '일자: ________________')

    add_para(doc, '')
    add_para(doc, '─── 문서 끝 ───')

    doc.save(out_path)
    print(f'  ✓ {out_path.name}')


# ════════════════════════════════════════════════════════════════════
# 메인
# ════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    out_dir = Path(__file__).resolve().parent.parent / 'docs'
    out_dir.mkdir(exist_ok=True)
    print(f'생성 디렉토리: {out_dir}')

    build_function_spec(out_dir / 'Pathwave_MVP_FunctionSpec_v1.0.docx')
    build_beacon_protocol(out_dir / 'Pathwave_MVP_BeaconProtocol_v1.0.docx')
    build_nda(out_dir / 'Pathwave_MVP_NDA_v1.0.docx')

    print('\n완료. 3 종 별첨 docx 생성.')
