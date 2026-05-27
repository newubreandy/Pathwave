"""PathWave 사업 추진 계획서 (창업지원단 제출용) → MS Word(.docx) 생성.

실행:
    cd /Users/m5pro16/Desktop/pathwave
    ./venv/bin/python docs/proposal/build_proposal.py

산출물:
    docs/proposal/PathWave_사업추진계획서_2026-05-03.docx

특징:
- 보고용 추상화 (오버커밋 방지)
- 9월 완료 일정 기준
- 사업비 = 개발비 + 하드웨어 + 소프트웨어/서버 + 인증 (분리)
- 선금 50% / 잔금 50% (9월말)
- 심의 수정 대응 무상
- 위험/유의 항목 빨간색 강조
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# ── 색상 ──────────────────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x4A, 0x23, 0xBF)
DARK_PURPLE = RGBColor(0x33, 0x18, 0x99)
GREY        = RGBColor(0x66, 0x66, 0x66)
DARK        = RGBColor(0x1A, 0x1A, 0x1A)
RED         = RGBColor(0xCC, 0x00, 0x00)
DARK_RED    = RGBColor(0x99, 0x00, 0x00)
LIGHT_BG    = "F4F0FF"
RED_BG      = "FFE5E5"


# ── 헬퍼 ──────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_korean_font(run, size=10, bold=False, color=None):
    run.font.name = 'Malgun Gothic'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_korean_font(run, size=28, bold=True, color=color or PURPLE)
    elif level == 1:
        run = p.add_run(text)
        set_korean_font(run, size=18, bold=True, color=color or PURPLE)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run = p.add_run(text)
        set_korean_font(run, size=13, bold=True, color=color or DARK_PURPLE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run = p.add_run(text)
        set_korean_font(run, size=11, bold=True, color=color or DARK)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, size=10, bold=False, color=None, indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_korean_font(run, size=size)
    return p


def add_bullet_red(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run('[유의] ')
    set_korean_font(run, size=size, bold=True, color=RED)
    run = p.add_run(text)
    set_korean_font(run, size=size, color=RED)
    return p


def add_numbered(doc, text, size=10):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    set_korean_font(run, size=size)


def add_table(doc, headers, rows, col_widths=None, header_bg=LIGHT_BG, font_size=10,
              red_rows=None):
    red_rows = red_rows or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_korean_font(run, size=font_size, bold=True, color=DARK_PURPLE)

    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        is_red_row = r in red_rows
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            color = RED if is_red_row else None
            set_korean_font(run, size=font_size, color=color, bold=is_red_row)
            if is_red_row:
                set_cell_bg(cell, RED_BG)
    return table


def add_note_box(doc, text, color=RGBColor(0x66, 0x46, 0x00), bg="FFFBE6", icon='[참고] '):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(icon + text)
    set_korean_font(run, size=10, color=color)
    return table


def add_red_box(doc, text):
    return add_note_box(doc, text, color=DARK_RED, bg=RED_BG, icon='[유의] ')


def add_page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  문서 시작
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


# ── 표지 ────────────────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PathWave')
set_korean_font(run, size=36, bold=True, color=PURPLE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('사업 추진 계획서')
set_korean_font(run, size=22, bold=True, color=DARK_PURPLE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 창업지원단 제출용 —')
set_korean_font(run, size=13, color=GREY)

for _ in range(5):
    doc.add_paragraph()

cover_table = doc.add_table(rows=6, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_info = [
    ('사업명',       'PathWave (BLE 비콘 기반 매장 통합 서비스)'),
    ('수행기관',     'PathWave (개인사업자 / 법인 전환 예정)'),
    ('대표자',       '___________________________'),
    ('추진 기간',    '계약일 ~ 2026년 9월'),
    ('작성일',       '2026-05-03'),
    ('버전',         'v1.0 (제출용)'),
]
for r, (k, v) in enumerate(cover_info):
    row = cover_table.rows[r]
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(11)
    set_cell_bg(row.cells[0], "F7F5FF")
    row.cells[0].text = ''
    p = row.cells[0].paragraphs[0]
    run = p.add_run(k)
    set_korean_font(run, size=11, bold=True, color=DARK_PURPLE)
    row.cells[1].text = ''
    p = row.cells[1].paragraphs[0]
    run = p.add_run(v)
    set_korean_font(run, size=11)

add_page_break(doc)


# ── 목차 ────────────────────────────────────────────────────────────────────
add_heading(doc, '목   차', level=1)
for item in [
    '1. 사업 개요',
    '2. 서비스 정의 및 사용자 구성',
    '3. 핵심 기능 (페르소나별)',
    '4. 시스템 구성',
    '5. 기술적 유의 사항',
    '6. 추진 일정',
    '7. 사업비 계획',
    '8. 산출물 및 평가 기준',
    '9. 운영 사항',
]:
    add_para(doc, item, size=11, indent=0.5)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. 사업 개요
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. 사업 개요', level=1)

add_heading(doc, '1.1 사업 목표', level=2)
add_para(doc,
    '오프라인 매장에 설치된 BLE 비콘을 매개로, 일반 회원의 매장 방문이 자동 인식되어 '
    'WiFi 연결 · 스탬프 적립 · 쿠폰 발급이 자연스럽게 이루어지는 매장-회원 통합 플랫폼을 구축한다. '
    '시설 사장님은 매장 운영 도구를, PathWave 운영자는 비콘 자산 관리 도구를 함께 갖춤으로써 '
    '실외 광고 의존도를 낮추고 단골 고객 관계를 데이터로 관리할 수 있는 환경을 제공한다.'
)

add_heading(doc, '1.2 추진 기간', level=2)
add_bullet(doc, '계약 체결일 ~ 2026년 9월 (테스트 완료 + 앱 심의 신청까지)')
add_bullet(doc, '단계별 마일스톤은 6장 참조')

add_heading(doc, '1.3 기대 효과', level=2)
for item in [
    '소상공인의 오프라인 광고 비용을 절감하면서 단골 데이터를 자체 보유',
    '회원 추천(와이파이 초대) 기반 자연스러운 사용자 확장 구조',
    'BLE 비콘 인프라의 유지·보수를 PathWave가 일괄 관리하여 매장 진입 장벽을 낮춤',
]:
    add_bullet(doc, item)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. 서비스 정의 및 사용자 구성
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. 서비스 정의 및 사용자 구성', level=1)

add_heading(doc, '2.1 서비스 한 줄 정의', level=2)
add_para(doc,
    '“매장에 들어서면 알아서 연결되는 통합 회원제 서비스” — BLE 비콘이 사용자를 인식하여 '
    'WiFi · 스탬프 · 쿠폰 · 알림이 자동으로 동작한다.'
)

add_heading(doc, '2.2 사용자 구성 (4 페르소나)', level=2)
add_table(doc,
    headers=['페르소나', '핵심 행위', '접근 채널'],
    rows=[
        ('일반 회원 (User)', '매장 방문 시 자동 인식 → WiFi/스탬프/쿠폰, 와이파이 초대 발송', '모바일 앱'),
        ('시설 사장님 (Owner)', '매장·비콘 등록, 스탬프/쿠폰/알림 운영, 직원 관리, 리포트 확인', '반응형 웹'),
        ('직원 (Staff)', '매장 일상 업무 (쿠폰 사용 처리, 고객 응대) — 권한 한정', '반응형 웹'),
        ('운영자 (Super Admin)', '비콘 입고·할당·재고·배터리 모니터링, 사장 가입 승인, 결제·정산', '관리자 웹'),
    ],
    col_widths=[4.5, 9, 3],
)

add_heading(doc, '2.3 가입 및 이용 흐름 (회원 폐쇄형)', level=2)
add_para(doc,
    'PathWave는 회원 전용 서비스이며 비회원은 기능을 이용할 수 없다. '
    '회원 확장은 기존 회원이 매장에서 동행한 지인에게 “와이파이 초대 링크”를 전달하면, '
    '지인이 회원가입 절차를 거친 뒤 동일 서비스를 이용할 수 있는 방식으로 이루어진다.'
)
add_table(doc,
    headers=['단계', '주체', '동작'],
    rows=[
        ('① 초대',    '기존 회원',  '매장에서 “와이파이 초대” 메뉴를 통해 지인에게 가입 링크 전달 (카카오톡/문자 등)'),
        ('② 가입',    '지인',       '링크 진입 → 약관 동의 → 회원가입 (이메일·휴대폰 인증)'),
        ('③ 자동 연결', '신규 회원', '같은 매장 BLE 비콘 인식 → WiFi 연결 + 첫 스탬프 적립'),
        ('④ 보상',    '초대한 회원', '초대 보상 쿠폰/스탬프 자동 발급 (정책 별도)'),
    ],
    col_widths=[2.5, 3, 10.5],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. 핵심 기능 (페르소나별)
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, '3. 핵심 기능', level=1)

add_heading(doc, '3.1 일반 회원 (User)', level=2)
for item in [
    '회원가입 / 로그인 (이메일·휴대폰 인증)',
    '소셜 로그인 (Google, Apple, 카카오 등)',
    '매장 검색 및 즐겨찾기',
    '매장 방문 자동 인식 (BLE 비콘) → WiFi 연결',
    '스탬프 자동 적립 / 쿠폰 자동 발급',
    '쿠폰 사용 (매장 직원 처리)',
    '매장과 1:1 채팅',
    '푸시 알림',
    '와이파이 초대 (지인 가입 링크 발송 + 보상)',
    '계정/알림 설정, 회원 탈퇴',
]:
    add_bullet(doc, item)

add_heading(doc, '3.2 시설 사장님 (Owner)', level=2)
add_para(doc, '가입 시 사업자 진위 확인 + 본인 이메일 + 휴대폰 인증으로 본인 검증을 수행한다.', indent=0.3)
for item in [
    '사장 회원가입 및 운영자 승인 절차',
    '매장 정보 등록 (이름·주소·영업시간·이미지·다국어)',
    '비콘 클레임 (지급받은 비콘 SN으로 본인 매장에 연결)',
    '스탬프 서비스 신청·등록·수정·삭제',
    '쿠폰 서비스 신청·등록·수정·삭제',
    '알림(푸시) 서비스 신청·등록·수정·삭제',
    '직원 초대 및 권한 관리',
    '매출·스탬프·쿠폰 리포트 확인',
    '결제·구독 관리',
]:
    add_bullet(doc, item)

add_heading(doc, '3.3 직원 (Staff)', level=2)
add_para(doc, '사장이 부여한 권한 범위 내에서 매장 일상 업무를 처리한다.', indent=0.3)
for item in [
    '직원 초대 수락 후 로그인',
    '매장 정보 조회 (수정은 권한별)',
    '쿠폰 사용 처리, 스탬프 수동 적립',
    '고객 채팅 응대',
    '리포트 조회 (열람 권한 한정)',
]:
    add_bullet(doc, item)

add_heading(doc, '3.4 운영자 (Super Admin)', level=2)
for item in [
    '운영자 로그인 (보안 강화)',
    '비콘 자산 관리 — 등록 / 일괄 입고 / 할당 / 회수 / 폐기',
    '비콘 재고 현황 및 매장별 분포 조회',
    '비콘 배터리 모니터링 (저전력 알람, 교체 일정 관리)',
    '사장 가입 승인 (사업자등록증 검토)',
    '계정 정지 / 재활성화',
    '결제·환불·구독 관리',
    '시스템 공지 (전체/매장/회원)',
]:
    add_bullet(doc, item)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. 시스템 구성
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. 시스템 구성', level=1)
add_table(doc,
    headers=['구성 요소', '주요 역할', '플랫폼'],
    rows=[
        ('① 백엔드 API 서버',         '인증·시설·비콘·스탬프·쿠폰·결제·푸시·채팅 처리',  'REST + SSE'),
        ('② 시설관리자 반응형 웹',    '사장님/직원의 매장 운영 도구',                    'PC + 모바일'),
        ('③ 사용자 모바일 앱',        '회원의 매장 이용 도구, BLE 자동 인식',            'iOS + Android'),
        ('④ 운영자 콘솔',             '비콘 자산 관리, 가입 승인, 결제·정산',            '웹 (PC)'),
        ('⑤ 인프라',                  '클라우드 호스팅, 배포, 모니터링, 백업',            '클라우드'),
    ],
    col_widths=[5, 8, 3],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. 기술적 유의 사항 ⚠
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, '5. 기술적 유의 사항 (사전 식별)', level=1, color=DARK_RED)
add_red_box(doc,
    '본 장은 사업 진행 중 일정·비용에 영향을 줄 수 있는 외부 의존 및 플랫폼 제약을 사전 식별한 것이다. '
    '미인지 시 후속 평가 단계에서 보완이 필요할 수 있어 진행 초기에 분명히 명시한다.'
)

add_heading(doc, '5.1 모바일 OS 제약', level=2, color=DARK_RED)
add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ('iOS / Android WiFi 연결',
         'OS 정책상 “완전 자동 연결”은 불가하며, 사용자의 1회 동의 후 자동 재연결 방식으로 구현한다.'),
        ('iOS Background BLE 스캔',
         '일반 BLE 스캔은 백그라운드에서 제약. 비콘 영역 모니터링 (iBeacon/CoreLocation)으로 보완.'),
        ('Android 12+ BLE 권한',
         'BLUETOOTH_SCAN/CONNECT 런타임 권한, 신규 권한 모델 적용.'),
    ],
    col_widths=[5, 11],
    red_rows=[0, 1, 2],
)

add_heading(doc, '5.2 외부 서비스 의존', level=2, color=DARK_RED)
add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ('카카오 / 네이버 소셜 로그인',
         '국내 소셜 로그인은 Firebase Authentication 직접 미지원으로, 별도 토큰 변환 흐름 필요.'),
        ('PG (전자결제) 연동',
         '법인 전환 후 사업자 심사를 거쳐 실 키 발급. 그 이전엔 sandbox 단계로 진행.'),
        ('FCM 푸시 알림',
         'iOS는 Apple Developer Program 가입 및 APNs 인증서 발급 필요.'),
        ('BLE 비콘 펌웨어',
         '비콘 제조사로부터 핸드셰이크 프로토콜 사양 확보 후 진행.'),
    ],
    col_widths=[5, 11],
    red_rows=[0, 1, 2, 3],
)

add_heading(doc, '5.3 보안 및 규제', level=2, color=DARK_RED)
add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ('개인정보보호법 / 정보통신망법',
         '회원가입·결제 화면의 개인정보 처리방침 노출, 만 14세 미만 가입 제한, 파기 절차 준수.'),
        ('WiFi 비밀번호 보호',
         '서버 저장 시 강암호화(AES-256-GCM 등), 클라이언트 전송은 HTTPS 위에서만 허용.'),
        ('앱 심의',
         '구글/애플 스토어 심의 기준에 맞춰 약관·개인정보 동의 흐름 정비.'),
    ],
    col_widths=[5, 11],
    red_rows=[0, 1, 2],
)

add_heading(doc, '5.4 운영 의존성', level=2, color=DARK_RED)
add_bullet_red(doc, '법인 등록 이후 발급 가능한 외부 자원(PG·APNs·일부 SDK) 일정에 본 사업 일정이 일부 종속됨')
add_bullet_red(doc, '베타 테스트 단계까지는 sandbox/stub 모드로 진행, 실 키 교체는 정식 런칭 직전 1~2주에 수행')


# ═══════════════════════════════════════════════════════════════════════════════
# 6. 추진 일정
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, '6. 추진 일정', level=1)
add_table(doc,
    headers=['단계', '내용', '기간(예시)'],
    rows=[
        ('M0', '킥오프, 자료 정비, 아키텍처 확정', '계약 후 2주'),
        ('M1', '백엔드 코어 기능 구현 (인증·시설·비콘·스탬프)', '약 6주'),
        ('M2', '백엔드 확장 기능 (쿠폰·알림·푸시·채팅·결제·리포트)', '약 4주'),
        ('M3', '시설관리자 웹 + 운영자 콘솔', '약 6주'),
        ('M4', '사용자 모바일 앱 + BLE 통합', '약 8주'),
        ('M5', '인프라 구성 + 통합 검수', '약 3주'),
        ('M6', '베타 테스트 + 앱 심의 신청', '~ 2026년 9월'),
    ],
    col_widths=[2, 11, 3],
)
add_red_box(doc,
    '내부적으로 조기 완료가 가능할 경우 사전 베타 오픈 또는 정식 런칭 시점을 앞당길 수 있다. '
    '단, 본 계획서의 공식 일정은 9월 완료 기준으로 한다.'
)


# ═══════════════════════════════════════════════════════════════════════════════
# 7. 사업비 계획
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. 사업비 계획', level=1)
add_para(doc,
    '본 사업비는 (가) 개발 인건비 및 일반비를 포함하며, '
    '(나) 하드웨어, (다) 소프트웨어/서버, (라) 인증 비용은 별도 항목으로 산정한다.'
)

add_heading(doc, '7.1 개발 인건비 및 일반비', level=2)
add_table(doc,
    headers=['항목', '내용', '금액 (VAT 별도)'],
    rows=[
        ('인건비',     '백엔드·웹·모바일 통합 개발 (수행기관 자체 인력)', '　원'),
        ('일반비',     '회의비, 출장비, 도서, 사무용품 등',              '　원'),
        ('소계',       '—',                                              '　원'),
    ],
    col_widths=[4, 9, 4],
)

add_heading(doc, '7.2 하드웨어 비용 (별도)', level=2)
add_table(doc,
    headers=['항목', '내용', '금액 (VAT 별도)'],
    rows=[
        ('BLE 비콘 자재', 'FSC-BP108B 등 비콘 본체 + 부착재 (수량 협의)', '　원'),
        ('테스트 단말',   'iOS / Android 검증용 단말 (대여 또는 구매)',  '　원'),
        ('소계',          '—',                                            '　원'),
    ],
    col_widths=[4, 9, 4],
)

add_heading(doc, '7.3 소프트웨어 / 서버 비용 (별도)', level=2)
add_table(doc,
    headers=['항목', '내용', '금액 (VAT 별도)'],
    rows=[
        ('클라우드 호스팅',     '백엔드 서버 / DB / 정적 호스팅 (월 단위)',         '　원/월'),
        ('Firebase (FCM)',      '푸시 알림 서비스, 일정 사용량 무료 + 초과분',      '　원'),
        ('Google Translate API', '다국어 번역 (사용량 기반 유료)',                  '　원'),
        ('PG 수수료',           '실 결제 수수료 (거래액 기반, 정식 런칭 후 발생)',  '거래액 기반'),
        ('도메인 / SSL',        '도메인 등록 + HTTPS 인증서 (연 단위)',             '　원/년'),
        ('소계',                '—',                                                '　원'),
    ],
    col_widths=[4, 9, 4],
)

add_heading(doc, '7.4 인증 / 심의 비용 (별도)', level=2)
add_table(doc,
    headers=['항목', '내용', '금액 (VAT 별도)'],
    rows=[
        ('Apple Developer Program',  '연 $99 (한화 환산)',                       '　원/년'),
        ('Google Play Console',       '최초 $25 (1회)',                          '　원'),
        ('앱 심의 대응',              '심의 거절 시 보완·재심의 (본 계약 무상)', '無'),
        ('소계',                      '—',                                       '　원'),
    ],
    col_widths=[4, 9, 4],
)

add_heading(doc, '7.5 지급 조건', level=2)
add_table(doc,
    headers=['구분', '비율', '지급 시기'],
    rows=[
        ('선금', '50%', '계약 체결 시'),
        ('잔금', '50%', '2026년 9월 말 (최종 산출물 제출 후)'),
    ],
    col_widths=[4, 4, 9],
)

add_heading(doc, '7.6 사후 대응', level=2)
add_para(doc,
    '앱 심의 결과에 따른 보완 작업 및 평가 단계의 수정 요청은 본 계약 범위 내에서 무상으로 대응한다. '
    '단, 사업 범위 자체의 변경(신규 기능 추가 등)은 별도 협의 후 추가 비용으로 처리한다.'
)


# ═══════════════════════════════════════════════════════════════════════════════
# 8. 산출물 및 평가 기준
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, '8. 산출물 및 평가 기준', level=1)

add_heading(doc, '8.1 최종 산출물', level=2)
for item in [
    '백엔드 시스템 (운영 환경 배포)',
    '시설관리자 반응형 웹 (운영 도메인 호스팅)',
    '사용자 모바일 앱 (iOS/Android, 스토어 심의 신청 완료 상태)',
    '운영자 콘솔 (사내 관리용 호스팅)',
    '운영 매뉴얼 및 사용자 가이드',
    '소스 코드 일체 (수행기관 보유)',
]:
    add_bullet(doc, item)

add_heading(doc, '8.2 평가 기준', level=2)
add_table(doc,
    headers=['구분', '기준'],
    rows=[
        ('기능',       '3장의 페르소나별 핵심 기능 동작 시연 가능'),
        ('품질',       '주요 사용자 흐름 (가입·BLE 자동 인식·스탬프·쿠폰·결제) 무결점'),
        ('완성도',     '베타 테스트 단계 진입 가능, 앱 심의 신청 완료'),
        ('문서화',     '운영 매뉴얼 및 사용자 가이드 제출'),
    ],
    col_widths=[3, 13],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 9. 운영 사항
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. 운영 사항', level=1)
for item in [
    '미팅 및 업무 교류는 정기 일정 없이 필요시 상시 진행한다.',
    '주요 의사결정은 이메일 또는 메신저로 기록을 남긴다.',
    '소스 코드 및 산출물의 모든 권리는 수행기관에 귀속한다.',
    '본 사업 진행 중 알게 된 평가기관·수행기관 양측의 정보는 외부 유출 금지.',
    '사업 범위의 중대 변경은 서면 합의 후 적용한다.',
    '평가 단계에서 추가 보완 요청 시 본 계약 범위 내 무상 대응을 원칙으로 하되, 신규 기능 추가는 별도 협의.',
]:
    add_numbered(doc, item)


# ── 서명란 ──────────────────────────────────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()

sig_table = doc.add_table(rows=1, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

left = sig_table.rows[0].cells[0]
left.text = ''
p = left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('수행기관\n')
set_korean_font(run, size=12, bold=True, color=DARK_PURPLE)
run = p.add_run('PathWave\n\n\n')
set_korean_font(run, size=10)
run = p.add_run('대표: ___________________ (인)\n\n')
set_korean_font(run, size=10)
run = p.add_run('날짜: 2026년 ___월 ___일')
set_korean_font(run, size=10)

right = sig_table.rows[0].cells[1]
right.text = ''
p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('평가기관\n')
set_korean_font(run, size=12, bold=True, color=DARK_PURPLE)
run = p.add_run('___________________________\n\n\n')
set_korean_font(run, size=10)
run = p.add_run('담당: ___________________ (인)\n\n')
set_korean_font(run, size=10)
run = p.add_run('날짜: 2026년 ___월 ___일')
set_korean_font(run, size=10)


# ── 푸터 ────────────────────────────────────────────────────────────────────
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— PathWave 사업 추진 계획서 v1.0 (창업지원단 제출용) —')
set_korean_font(run, size=9, color=GREY)


# ── 저장 ────────────────────────────────────────────────────────────────────
out_path = Path('/Users/m5pro16/Desktop/pathwave/docs/proposal/PathWave_사업추진계획서_2026-05-03.docx')
doc.save(out_path)
print(f'✅ 저장 완료: {out_path}')
print(f'   크기: {out_path.stat().st_size:,} bytes')
